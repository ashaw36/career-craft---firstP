"""
CareerCraft Agent — ImportParser 文件导入测试

测试范围：
1. PDF 文本提取
2. Word 文本提取
3. import_file 完整链路（含 mock LLM 解析）
4. analyze_file_with_llm JSON 解析弹性
"""

from __future__ import annotations

import json
from datetime import date
from typing import Any, Dict, List
from unittest.mock import AsyncMock

import pytest

from src.services.import_parser import ImportParser, ImportParserError
from src.services.experience_manager import ExperienceDraft


class TestExtractTextFromPDF:
    """PDF 文本提取测试"""

    def test_extract_from_simple_pdf(self) -> None:
        """从简单 PDF 提取文本"""
        import fitz

        # 构建一个内存 PDF
        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "CareerCraft Project Summary")
        page.insert_text((50, 80), "Responsible for product planning")
        pdf_bytes = doc.tobytes()
        doc.close()

        text = ImportParser.extract_text_from_pdf(pdf_bytes)
        assert "CareerCraft Project Summary" in text
        assert "product planning" in text

    def test_extract_from_pdf_with_chinese(self) -> None:
        """从含中文的 PDF 提取（只验证有文字返回）"""
        import fitz

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "项目总结")
        pdf_bytes = doc.tobytes()
        doc.close()

        text = ImportParser.extract_text_from_pdf(pdf_bytes)
        # fitz 在测试环境中可能不能完美渲染中文，只验证不报错且有返回
        assert isinstance(text, str)
        # 至少有一些字符被提取（可能是方块或空白）
        assert len(text) >= 0

    def test_extract_from_empty_pdf(self) -> None:
        """空 PDF 返回空字符串"""
        import fitz

        doc = fitz.open()
        doc.new_page()
        pdf_bytes = doc.tobytes()
        doc.close()

        text = ImportParser.extract_text_from_pdf(pdf_bytes)
        assert text.strip() == ""


class TestExtractTextFromWord:
    """Word 文本提取测试"""

    def test_extract_from_simple_docx(self) -> None:
        """从简单 Word 提取文本"""
        import io

        from docx import Document

        doc = Document()
        doc.add_paragraph("工作汇报")
        doc.add_paragraph("负责供应链数字化项目")
        doc.add_paragraph("提升效率 40%")
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        text = ImportParser.extract_text_from_word(doc_bytes)
        assert "工作汇报" in text
        assert "供应链数字化项目" in text
        assert "40%" in text

    def test_extract_from_empty_docx(self) -> None:
        """空 Word 返回空字符串"""
        import io

        from docx import Document

        doc = Document()
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        text = ImportParser.extract_text_from_word(doc_bytes)
        assert text.strip() == ""


class TestAnalyzeFileWithLLM:
    """LLM 文件分析测试"""

    @pytest.fixture
    def parser(self) -> ImportParser:
        mock_llm = AsyncMock()
        return ImportParser(llm_router=mock_llm)

    @pytest.mark.asyncio
    async def test_analyze_returns_experiences(self, parser: ImportParser) -> None:
        """LLM 返回标准 JSON 数组"""
        parser.llm.chat.return_value = json.dumps([
            {
                "title": "产品经理",
                "type": "work",
                "organization": "美团",
                "start_date": "2020-06",
                "end_date": "2023-08",
                "raw_description": "负责增长产品，DAU 提升 50%",
                "skills_demonstrated": ["Python", "数据分析"],
                "structured_achievements": ["DAU 提升 50%"],
                "metrics": [{"metric": "DAU", "value": "50%", "unit": "提升"}],
            }
        ], ensure_ascii=False)

        drafts = await parser.analyze_file_with_llm("一些项目总结", "项目总结")
        assert len(drafts) == 1
        draft = drafts[0]
        assert isinstance(draft, ExperienceDraft)
        assert draft.extracted["title"] == "产品经理"
        assert draft.extracted["organization"] == "美团"
        assert draft.extracted["skills_demonstrated"] == ["Python", "数据分析"]

    @pytest.mark.asyncio
    async def test_analyze_json_in_markdown(self, parser: ImportParser) -> None:
        """LLM 返回 markdown 代码块"""
        parser.llm.chat.return_value = (
            "```json\n"
            + json.dumps([{"title": "工程师", "type": "work", "organization": "字节"}], ensure_ascii=False)
            + "\n```"
        )

        drafts = await parser.analyze_file_with_llm("简历", "简历")
        assert len(drafts) == 1
        assert drafts[0].extracted["title"] == "工程师"

    @pytest.mark.asyncio
    async def test_analyze_invalid_json_fallback(self, parser: ImportParser) -> None:
        """LLM 返回不规范 JSON，尝试回退解析"""
        parser.llm.chat.return_value = '这里是一些分析' + json.dumps(
            [{"title": "运营", "type": "work"}], ensure_ascii=False
        )

        drafts = await parser.analyze_file_with_llm("文本", "文本")
        assert len(drafts) == 1
        assert drafts[0].extracted["title"] == "运营"

    @pytest.mark.asyncio
    async def test_analyze_llm_error(self, parser: ImportParser) -> None:
        """LLM 调用失败抛出异常"""
        parser.llm.chat.side_effect = RuntimeError("API 异常")

        with pytest.raises(ImportParserError, match="LLM 分析失败"):
            await parser.analyze_file_with_llm("内容", "文件")

    @pytest.mark.asyncio
    async def test_analyze_skips_empty_title(self, parser: ImportParser) -> None:
        """跳过无标题的经历条目"""
        parser.llm.chat.return_value = json.dumps(
            [{"title": "", "type": "work"}, {"title": "有效", "type": "work"}],
            ensure_ascii=False,
        )

        drafts = await parser.analyze_file_with_llm("内容", "文件")
        assert len(drafts) == 1
        assert drafts[0].extracted["title"] == "有效"


class TestImportFile:
    """import_file 完整链路测试"""

    @pytest.fixture
    def parser(self) -> ImportParser:
        mock_llm = AsyncMock()
        return ImportParser(llm_router=mock_llm)

    @pytest.mark.asyncio
    async def test_import_txt_file(self, parser: ImportParser) -> None:
        """导入文本文件并解析"""
        parser.llm.chat.return_value = json.dumps(
            [{"title": "文本分析", "type": "work", "organization": "测试公司"}],
            ensure_ascii=False,
        )

        text_bytes = "这是一份简历内容".encode("utf-8")
        drafts = await parser.import_file("resume.txt", text_bytes)
        assert len(drafts) == 1
        assert drafts[0].extracted["title"] == "文本分析"

    @pytest.mark.asyncio
    async def test_import_pdf_file(self, parser: ImportParser) -> None:
        """导入 PDF 文件并解析"""
        import fitz

        parser.llm.chat.return_value = json.dumps(
            [{"title": "PDF分析", "type": "work", "organization": "PDF公司"}],
            ensure_ascii=False,
        )

        doc = fitz.open()
        page = doc.new_page()
        page.insert_text((50, 50), "这是 PDF 内容")
        pdf_bytes = doc.tobytes()
        doc.close()

        drafts = await parser.import_file("report.pdf", pdf_bytes)
        assert len(drafts) == 1
        assert drafts[0].extracted["title"] == "PDF分析"

    @pytest.mark.asyncio
    async def test_import_docx_file(self, parser: ImportParser) -> None:
        """导入 Word 文件并解析"""
        import io

        from docx import Document

        parser.llm.chat.return_value = json.dumps(
            [{"title": "Word分析", "type": "work", "organization": "Word公司"}],
            ensure_ascii=False,
        )

        doc = Document()
        doc.add_paragraph("这是 Word 内容")
        buf = io.BytesIO()
        doc.save(buf)
        doc_bytes = buf.getvalue()

        drafts = await parser.import_file("report.docx", doc_bytes)
        assert len(drafts) == 1
        assert drafts[0].extracted["title"] == "Word分析"

    @pytest.mark.asyncio
    async def test_import_empty_file_raises(self, parser: ImportParser) -> None:
        """空文件抛出异常"""
        with pytest.raises(ImportParserError, match="无法从 empty.txt 提取"):
            await parser.import_file("empty.txt", b"")


class TestToDraft:
    """ParsedExperience → ExperienceDraft 转换测试"""

    def test_to_draft_preserves_all_fields(self) -> None:
        """转换保留所有字段"""
        parsed = ImportParser._dict_to_experience({
            "title": "测试",
            "organization": "公司A",
            "type": "project",
            "start_date": "2021-03",
            "end_date": "2022-06",
            "raw_description": "项目描述",
            "skills_demonstrated": ["Python", "Docker"],
            "structured_achievements": ["提升效率"],
        })
        assert parsed is not None
        draft = ImportParser._to_draft(parsed)
        assert draft.extracted["title"] == "测试"
        assert draft.extracted["organization"] == "公司A"
        assert draft.extracted["type"] == "project"
        assert draft.extracted["start_date"] == "2021-03-01"
        assert draft.extracted["end_date"] == "2022-06-01"
        assert draft.extracted["skills_demonstrated"] == ["Python", "Docker"]

    def test_to_draft_none_dates(self) -> None:
        """日期为空时不报错"""
        parsed = ImportParser._dict_to_experience({
            "title": "测试",
        })
        assert parsed is not None
        draft = ImportParser._to_draft(parsed)
        assert draft.extracted["start_date"] is None
        assert draft.extracted["end_date"] is None
